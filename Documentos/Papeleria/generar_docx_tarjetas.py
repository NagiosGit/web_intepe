import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

output_dir = r"c:\Users\WINTEPE\Desktop\PROYECTOS_IA\Web_Intepe\Documentos\Papeleria"
public_dir = r"c:\Users\WINTEPE\Desktop\PROYECTOS_IA\Web_Intepe\public\downloads"

def set_cell_margins(cell, top=60, bottom=60, left=120, right=120):
    """Set cell padding in dxa (1 pt = 20 dxa, 1 mm = 56.7 dxa)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_border(cell, color="CBD5E1", sz="6", val="single"):
    """Set outer border around cell"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        </w:tcBorders>
    ''')
    tcPr.append(tcBorders)

def set_no_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="none"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
            <w:insideH w:val="none"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(tblBorders)

def build_card_docx(profile, qr_img_path, filename):
    doc = Document()
    
    # Page Setup: Standard Letter (8.5 x 11.0 in)
    # Margins: Top 0.3 in, Bottom 0.3 in, Left 0.65 in, Right 0.65 in
    # Usable height: 10.4 in (264 mm). 4 rows * 2.15 in = 8.6 in -> Fits easily on 1 page!
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        section.top_margin = Inches(0.28)
        section.bottom_margin = Inches(0.28)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)
        
    card_w = Inches(3.35)  # 85 mm
    card_h = Inches(2.15)  # 54.6 mm
    
    # =========================================================================
    # PÁGINA 1: PLIEGO FRENTE (8 TARJETAS: 4 FILAS X 2 COLUMNAS)
    # =========================================================================
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_header.paragraph_format.space_before = Pt(0)
    p_header.paragraph_format.space_after = Pt(2)
    p_header.paragraph_format.line_spacing = Pt(10)
    r_h = p_header.add_run("📄 HOJA 1 DE 2: PLIEGO FRENTE (8 TARJETAS - ANVERSO)")
    r_h.bold = True
    r_h.font.name = "Arial"
    r_h.font.size = Pt(8)
    r_h.font.color.rgb = RGBColor(16, 185, 129)
    
    table_front = doc.add_table(rows=4, cols=2)
    table_front.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row in table_front.rows:
        row.height = card_h
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for cell in row.cells:
            cell.width = card_w
            set_cell_margins(cell, top=70, bottom=60, left=130, right=130)
            set_cell_border(cell, color="CBD5E1", sz="6")
            
            # Line 1: Header Brand (Left) + Tag (Right via Tab Stop)
            p1 = cell.paragraphs[0]
            p1.paragraph_format.space_before = Pt(0)
            p1.paragraph_format.space_after = Pt(0)
            p1.paragraph_format.line_spacing = Pt(12)
            p1.paragraph_format.tab_stops.add_tab_stop(Inches(3.1), WD_TAB_ALIGNMENT.RIGHT)
            
            r_b = p1.add_run("INTEPE ")
            r_b.bold = True
            r_b.font.name = "Arial"
            r_b.font.size = Pt(12.5)
            r_b.font.color.rgb = RGBColor(15, 23, 42)
            
            r_sas = p1.add_run("S.A.S.")
            r_sas.bold = True
            r_sas.font.name = "Arial"
            r_sas.font.size = Pt(12.5)
            r_sas.font.color.rgb = RGBColor(255, 113, 32)
            
            p1.add_run("\t")
            r_tag = p1.add_run(f"[{profile['tag']}]")
            r_tag.bold = True
            r_tag.font.name = "Arial"
            r_tag.font.size = Pt(8)
            r_tag.font.color.rgb = RGBColor(194, 65, 12)
            
            # Line 2: Razón Social
            p2 = cell.add_paragraph()
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(5)
            p2.paragraph_format.line_spacing = Pt(9)
            r_sub = p2.add_run("Informática y Tecnología Penagos S.A.S.")
            r_sub.font.name = "Arial"
            r_sub.font.size = Pt(7.5)
            r_sub.font.color.rgb = RGBColor(100, 116, 139)
            
            # Line 3: Name
            p3 = cell.add_paragraph()
            p3.paragraph_format.space_before = Pt(0)
            p3.paragraph_format.space_after = Pt(0)
            p3.paragraph_format.line_spacing = Pt(14)
            r_name = p3.add_run(profile['name'])
            r_name.bold = True
            r_name.font.name = "Arial"
            r_name.font.size = Pt(13.5)
            r_name.font.color.rgb = RGBColor(15, 23, 42)
            
            # Line 4: Title & Department
            p4 = cell.add_paragraph()
            p4.paragraph_format.space_before = Pt(0)
            p4.paragraph_format.space_after = Pt(6)
            p4.paragraph_format.line_spacing = Pt(10)
            r_title = p4.add_run(profile['title'] + "\n")
            r_title.bold = True
            r_title.font.name = "Arial"
            r_title.font.size = Pt(9.5)
            r_title.font.color.rgb = RGBColor(234, 88, 12)
            
            r_dept = p4.add_run(profile['department'])
            r_dept.font.name = "Arial"
            r_dept.font.size = Pt(7.5)
            r_dept.font.color.rgb = RGBColor(100, 116, 139)
            
            # Line 5: Phone (Left) + Email (Right via Tab Stop)
            p5 = cell.add_paragraph()
            p5.paragraph_format.space_before = Pt(0)
            p5.paragraph_format.space_after = Pt(1)
            p5.paragraph_format.line_spacing = Pt(10)
            p5.paragraph_format.tab_stops.add_tab_stop(Inches(3.1), WD_TAB_ALIGNMENT.RIGHT)
            
            r_ph = p5.add_run(f"📞 {profile['mobile']}")
            r_ph.bold = True
            r_ph.font.name = "Arial"
            r_ph.font.size = Pt(8)
            r_ph.font.color.rgb = RGBColor(30, 41, 59)
            
            p5.add_run("\t")
            r_em = p5.add_run(f"✉️ {profile['email']}")
            r_em.bold = True
            r_em.font.name = "Arial"
            r_em.font.size = Pt(8)
            r_em.font.color.rgb = RGBColor(30, 41, 59)
            
            # Line 6: Website & NIT
            p6 = cell.add_paragraph()
            p6.paragraph_format.space_before = Pt(0)
            p6.paragraph_format.space_after = Pt(0)
            p6.paragraph_format.line_spacing = Pt(9)
            r_w = p6.add_run(f"🌐 www.intepe.net  •  NIT: {profile['nit']}")
            r_w.font.name = "Arial"
            r_w.font.size = Pt(7.5)
            r_w.font.color.rgb = RGBColor(100, 116, 139)

    # Page Break for Page 2
    doc.add_page_break()
    
    # =========================================================================
    # PÁGINA 2: PLIEGO REVERSO (8 TARJETAS CON QR - DÚPLEX)
    # =========================================================================
    p_header2 = doc.add_paragraph()
    p_header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_header2.paragraph_format.space_before = Pt(0)
    p_header2.paragraph_format.space_after = Pt(2)
    p_header2.paragraph_format.line_spacing = Pt(10)
    r_h2 = p_header2.add_run("🔄 HOJA 2 DE 2: PLIEGO REVERSO (8 TARJETAS CON QR - DÚPLEX)")
    r_h2.bold = True
    r_h2.font.name = "Arial"
    r_h2.font.size = Pt(8)
    r_h2.font.color.rgb = RGBColor(14, 116, 144)
    
    table_back = doc.add_table(rows=4, cols=2)
    table_back.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row in table_back.rows:
        row.height = card_h
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for cell in row.cells:
            cell.width = card_w
            set_cell_margins(cell, top=70, bottom=60, left=130, right=130)
            set_cell_border(cell, color="CBD5E1", sz="6")
            
            # Line 1: Header Tag (Left) + City (Right via Tab Stop)
            p1 = cell.paragraphs[0]
            p1.paragraph_format.space_before = Pt(0)
            p1.paragraph_format.space_after = Pt(3)
            p1.paragraph_format.line_spacing = Pt(10)
            p1.paragraph_format.tab_stops.add_tab_stop(Inches(3.1), WD_TAB_ALIGNMENT.RIGHT)
            
            r_bt = p1.add_run(f"// {profile['back_tag']}")
            r_bt.bold = True
            r_bt.font.name = "Arial"
            r_bt.font.size = Pt(8)
            r_bt.font.color.rgb = RGBColor(194, 65, 12)
            
            p1.add_run("\t")
            r_ct = p1.add_run("Bogotá, CO")
            r_ct.font.name = "Arial"
            r_ct.font.size = Pt(7.5)
            r_ct.font.color.rgb = RGBColor(100, 116, 139)
            
            # Center: Table of 1 row x 2 cols for Services (Left) and QR Code (Right)
            t_mid = cell.add_table(rows=1, cols=2)
            set_no_borders(t_mid)
            t_mid.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            c_srv = t_mid.rows[0].cells[0]
            c_srv.width = Inches(2.2)
            set_cell_margins(c_srv, 0, 0, 0, 0)
            p_srv = c_srv.paragraphs[0]
            p_srv.paragraph_format.space_before = Pt(0)
            p_srv.paragraph_format.space_after = Pt(0)
            p_srv.paragraph_format.line_spacing = Pt(11)
            
            for s in profile['services']:
                r_s = p_srv.add_run(f"• {s}\n")
                r_s.bold = True
                r_s.font.name = "Arial"
                r_s.font.size = Pt(8)
                r_s.font.color.rgb = RGBColor(30, 41, 59)
                
            c_qr = t_mid.rows[0].cells[1]
            c_qr.width = Inches(0.95)
            set_cell_margins(c_qr, 0, 0, 0, 0)
            p_qr = c_qr.paragraphs[0]
            p_qr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_qr.paragraph_format.space_before = Pt(0)
            p_qr.paragraph_format.space_after = Pt(1)
            p_qr.paragraph_format.line_spacing = None
            
            if os.path.exists(qr_img_path):
                r_img = p_qr.add_run()
                r_img.add_picture(qr_img_path, width=Inches(0.68), height=Inches(0.68))
                
            p_qlbl = c_qr.add_paragraph()
            p_qlbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_qlbl.paragraph_format.space_before = Pt(0)
            p_qlbl.paragraph_format.space_after = Pt(0)
            p_qlbl.paragraph_format.line_spacing = Pt(8)
            r_ql = p_qlbl.add_run("ESCANEAR")
            r_ql.bold = True
            r_ql.font.name = "Arial"
            r_ql.font.size = Pt(6.5)
            r_ql.font.color.rgb = RGBColor(194, 65, 12)
            
            # Bottom: Address (Left) + intepe.net (Right via Tab Stop)
            p_bot = cell.add_paragraph()
            p_bot.paragraph_format.space_before = Pt(3)
            p_bot.paragraph_format.space_after = Pt(0)
            p_bot.paragraph_format.line_spacing = Pt(9)
            p_bot.paragraph_format.tab_stops.add_tab_stop(Inches(3.1), WD_TAB_ALIGNMENT.RIGHT)
            
            r_ad = p_bot.add_run(profile['address'])
            r_ad.font.name = "Arial"
            r_ad.font.size = Pt(7.5)
            r_ad.font.color.rgb = RGBColor(100, 116, 139)
            
            p_bot.add_run("\t")
            r_wb = p_bot.add_run("intepe.net")
            r_wb.bold = True
            r_wb.font.name = "Arial"
            r_wb.font.size = Pt(8.5)
            r_wb.font.color.rgb = RGBColor(255, 113, 32)
            
    out_file = os.path.join(output_dir, filename)
    doc.save(out_file)
    
    pub_file = os.path.join(public_dir, filename)
    doc.save(pub_file)
    print(f"Saved: {out_file} and {pub_file}")

profile_william = {
    'name': 'Ing. William Penagos',
    'title': 'Director General & Especialista TI',
    'department': 'Dirección de Operaciones & Soluciones',
    'mobile': '(+57) 313 386 2656',
    'email': 'director@intepe.net',
    'address': 'Cl. 152a # 55-44, Bogotá',
    'website': 'https://www.intepe.net/tarjeta?p=director',
    'nit': '830.066.815-0',
    'tag': 'TI SOLUTIONS',
    'back_tag': 'SERVICIOS TI CORPORATIVOS',
    'services': [
        'Outsourcing TI Integral',
        'Mesa de Ayuda L1/L2/L3',
        'Infraestructura & Servidores',
        'Desarrollo de Software a Medida'
    ]
}

profile_patricia = {
    'name': 'Patricia Muñoz',
    'title': 'Gerente de Recursos Humanos',
    'department': 'Gestión del Talento & Desarrollo Organizacional',
    'mobile': '(+57) 313 386 2656',
    'email': 'rrhh@intepe.net',
    'address': 'Cl. 152a # 55-44, Bogotá',
    'website': 'https://www.intepe.net/tarjeta?p=rrhh',
    'nit': '830.066.815-0',
    'tag': 'TALENTO TI',
    'back_tag': 'GESTIÓN HUMANA & TALENTO',
    'services': [
        'Gestión del Talento Humano',
        'Selección de Personal TI',
        'Bienestar & Clima Organizacional',
        'Capacitación & Desarrollo TI'
    ]
}

qr_william = os.path.join(output_dir, "qr_william.png")
qr_patricia = os.path.join(output_dir, "qr_patricia.png")

build_card_docx(profile_william, qr_william, "Pliego_Tarjetas_INTEPE_Doble_Cara.docx")
build_card_docx(profile_william, qr_william, "Pliego_Tarjetas_Ing_William_Penagos.docx")
build_card_docx(profile_patricia, qr_patricia, "Pliego_Tarjetas_Patricia_Munoz.docx")
print("All DOCX files regenerated with perfect 2-page fit!")
